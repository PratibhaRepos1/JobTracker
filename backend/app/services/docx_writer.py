"""Render tailored markdown into a .docx file.

Approach: tokenize with mistune, walk tokens, emit python-docx calls.
Handles headings (H1-H3), paragraphs with **bold** / *italic*, bullet/numbered lists,
horizontal rules, and links (rendered as plain text).
"""
from io import BytesIO
from typing import Iterable

import mistune
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor


H1_COLOR = RGBColor(0x1F, 0x4E, 0x79)
H2_COLOR = RGBColor(0x55, 0x55, 0x55)
BODY_FONT = "Calibri"
BODY_SIZE = Pt(11)


def _set_run_style(run, *, bold=False, italic=False, color=None, size=None, font=BODY_FONT):
    run.font.name = font
    run.font.size = size or BODY_SIZE
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color is not None:
        run.font.color.rgb = color


def _render_inline(paragraph, children: Iterable):
    """Render mistune inline tokens into a paragraph's runs."""
    for tok in children:
        t = tok.get("type")
        if t == "text":
            run = paragraph.add_run(tok.get("raw", ""))
            _set_run_style(run)
        elif t == "strong":
            run = paragraph.add_run(_collect_text(tok.get("children", [])))
            _set_run_style(run, bold=True)
        elif t == "emphasis":
            run = paragraph.add_run(_collect_text(tok.get("children", [])))
            _set_run_style(run, italic=True)
        elif t == "codespan":
            run = paragraph.add_run(tok.get("raw", ""))
            _set_run_style(run, font="Consolas")
        elif t == "link":
            run = paragraph.add_run(_collect_text(tok.get("children", [])))
            _set_run_style(run, color=RGBColor(0x1F, 0x4E, 0x79))
        elif t == "linebreak" or t == "softbreak":
            paragraph.add_run("\n")
        else:
            if "children" in tok:
                _render_inline(paragraph, tok["children"])
            elif "raw" in tok:
                _set_run_style(paragraph.add_run(tok["raw"]))


def _collect_text(children: Iterable) -> str:
    out = []
    for tok in children:
        if tok.get("type") in {"text", "codespan"}:
            out.append(tok.get("raw", ""))
        elif "children" in tok:
            out.append(_collect_text(tok["children"]))
        elif "raw" in tok:
            out.append(tok["raw"])
    return "".join(out)


def _add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        _set_run_style(run, bold=True, color=H1_COLOR, size=Pt(18))
    elif level == 2:
        _set_run_style(run, bold=True, color=H2_COLOR, size=Pt(13))
    else:
        _set_run_style(run, bold=True, size=Pt(11))


def _walk(doc: Document, tokens: Iterable):
    for tok in tokens:
        t = tok.get("type")
        if t == "heading":
            text = _collect_text(tok.get("children", []))
            _add_heading(doc, text, tok.get("attrs", {}).get("level", 2))
        elif t == "paragraph":
            p = doc.add_paragraph()
            _render_inline(p, tok.get("children", []))
        elif t == "list":
            ordered = tok.get("attrs", {}).get("ordered", False)
            style = "List Number" if ordered else "List Bullet"
            for item in tok.get("children", []):
                p = doc.add_paragraph(style=style)
                for child in item.get("children", []):
                    if child.get("type") == "block_text" or child.get("type") == "paragraph":
                        _render_inline(p, child.get("children", []))
                    elif child.get("type") == "list":
                        _walk(doc, [child])
        elif t == "thematic_break":
            doc.add_paragraph("─" * 40)
        elif t == "block_code":
            p = doc.add_paragraph()
            run = p.add_run(tok.get("raw", ""))
            _set_run_style(run, font="Consolas", size=Pt(10))
        elif t == "blank_line":
            continue
        else:
            if "children" in tok:
                _walk(doc, tok["children"])


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    md = mistune.create_markdown(renderer=None)
    tokens = md(markdown_text)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE

    _walk(doc, tokens)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
