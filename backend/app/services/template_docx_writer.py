"""Render tailored markdown into a .docx using the user's original Word file as a template.

We load the user's resume.docx, keep its document defaults (theme fonts, page setup,
headers/footers, named styles), clear the body, and re-emit tailored content with the
same run-level formatting rules: 16pt bold centered for the name, 12pt bold for section
headers, "List Paragraph" style for bullets, justified body, etc.

This produces output visually indistinguishable from hand-edited content in the original
template, without requiring the user to mark up the docx with placeholders.
"""
from io import BytesIO
from pathlib import Path
from typing import Iterable, Literal, Optional

import mistune
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

Mode = Literal["resume", "cover"]


NAME_SIZE_PT = 15        # H1 — the candidate's name
SUBTITLE_SIZE_PT = 11    # the line directly under the name
SECTION_SIZE_PT = 11.5   # H2 — section headers (Professional Summary, etc.)
SUBSECTION_SIZE_PT = 10.5  # H3 — role headers within Experience
BODY_SIZE_PT = 10        # body paragraphs and bullets

SECTION_SPACE_BEFORE_PT = 5.0
SECTION_SPACE_AFTER_PT = 1.5
BODY_LINE_SPACING = 1.0
BODY_SPACE_AFTER_PT = 1.5

PHOTO_WIDTH_INCHES = 0.95
LINK_COLOR_HEX = "0563C1"  # Word's default hyperlink theme color


def _clear_body_keep_sectpr(doc: Document) -> None:
    """Drop everything in the body except the final sectPr (margins/page/headers)."""
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sectPr:
            body.remove(child)


def _collect_text(children: Iterable) -> str:
    out: list[str] = []
    for tok in children:
        if tok.get("type") in {"text", "codespan"}:
            out.append(tok.get("raw", ""))
        elif "children" in tok:
            out.append(_collect_text(tok["children"]))
        elif "raw" in tok:
            out.append(tok["raw"])
    return "".join(out)


def _add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, size_pt: Optional[float] = None):
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    return run


def _add_hyperlink(paragraph, url: str, text: str, *, size_pt: Optional[float] = None) -> None:
    """Insert a real Word hyperlink (clickable, blue-underlined) into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_COLOR_HEX)
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    if size_pt is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size_pt * 2)))  # w:sz is half-points
        rPr.append(sz)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def _render_inline(paragraph, children: Iterable, *, base_size_pt: Optional[float] = None) -> None:
    """Walk mistune inline tokens, adding runs while honoring **bold** / *italic* / [text](url)."""
    for tok in children:
        t = tok.get("type")
        if t == "text":
            _add_run(paragraph, tok.get("raw", ""), size_pt=base_size_pt)
        elif t == "strong":
            _add_run(paragraph, _collect_text(tok.get("children", [])), bold=True, size_pt=base_size_pt)
        elif t == "emphasis":
            _add_run(paragraph, _collect_text(tok.get("children", [])), italic=True, size_pt=base_size_pt)
        elif t == "codespan":
            run = _add_run(paragraph, tok.get("raw", ""), size_pt=base_size_pt)
            run.font.name = "Consolas"
        elif t == "link":
            url = tok.get("attrs", {}).get("url") or tok.get("link", "")
            link_text = _collect_text(tok.get("children", [])) or url
            if url:
                _add_hyperlink(paragraph, url, link_text, size_pt=base_size_pt)
            else:
                _add_run(paragraph, link_text, size_pt=base_size_pt)
        elif t in ("linebreak", "softbreak"):
            _add_run(paragraph, " ", size_pt=base_size_pt)
        else:
            if "children" in tok:
                _render_inline(paragraph, tok["children"], base_size_pt=base_size_pt)


def _add_photo(doc: Document, photo_path: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(4.0)
    run = p.add_run()
    run.add_picture(str(photo_path), width=Inches(PHOTO_WIDTH_INCHES))


def _add_name(doc: Document, text: str, *, mode: Mode) -> None:
    p = doc.add_paragraph()
    if mode == "resume":
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        size = NAME_SIZE_PT
    else:
        # Cover letter: left-aligned, slightly larger
        size = 18.0
    p.paragraph_format.space_after = Pt(1.5)
    _add_run(p, text, bold=True, size_pt=size)


def _add_subtitle(doc: Document, children: Iterable, *, is_first_under_name: bool, mode: Mode) -> None:
    p = doc.add_paragraph()
    if mode == "resume":
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        size = SUBTITLE_SIZE_PT if is_first_under_name else None
    else:
        # Cover letter: left-aligned. Subtitle 12pt, contact lines 10pt.
        size = 12.0 if is_first_under_name else 10.0
    p.paragraph_format.space_after = Pt(2.0 if is_first_under_name else 1.0)
    _render_inline(p, children, base_size_pt=size)


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(SECTION_SPACE_BEFORE_PT)
    p.paragraph_format.space_after = Pt(SECTION_SPACE_AFTER_PT)
    _add_run(p, text, bold=True, size_pt=SECTION_SIZE_PT)


def _add_subsection_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4.0)
    p.paragraph_format.space_after = Pt(1.0)
    _add_run(p, text, bold=True, size_pt=SUBSECTION_SIZE_PT)


def _add_body_paragraph(
    doc: Document,
    children: Iterable,
    *,
    justify: bool = True,
    mode: Mode = "resume",
) -> None:
    p = doc.add_paragraph()
    if mode == "resume":
        if justify:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(BODY_SPACE_AFTER_PT)
        p.paragraph_format.line_spacing = BODY_LINE_SPACING
        _render_inline(p, children, base_size_pt=BODY_SIZE_PT)
    else:
        # Cover letter: left-aligned, slightly more breathing room between paragraphs.
        p.paragraph_format.space_after = Pt(6.0)
        p.paragraph_format.line_spacing = 1.15
        _render_inline(p, children, base_size_pt=11.0)


def _split_on_softbreaks(children: Iterable) -> list[list]:
    """Group inline tokens into segments separated by soft/hard linebreaks."""
    segments: list[list] = [[]]
    for c in children:
        if c.get("type") in ("softbreak", "linebreak"):
            segments.append([])
        else:
            segments[-1].append(c)
    return [s for s in segments if any(_collect_text([t]).strip() for t in s)]


def _add_bullet(doc: Document, children: Iterable, *, available_styles: set[str]) -> None:
    style_name = "List Paragraph" if "List Paragraph" in available_styles else "Normal"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_after = Pt(BODY_SPACE_AFTER_PT)
    p.paragraph_format.line_spacing = BODY_LINE_SPACING
    _add_run(p, "•  ", size_pt=BODY_SIZE_PT)
    _render_inline(p, children, base_size_pt=BODY_SIZE_PT)


def _set_margins(doc: Document, *, top: float, bottom: float, left: float, right: float) -> None:
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def markdown_to_docx_with_template(
    markdown_text: str,
    template_path: Path,
    *,
    photo_path: Optional[Path] = None,
    mode: Mode = "resume",
) -> bytes:
    doc = Document(str(template_path))
    available_styles = {s.name for s in doc.styles}
    _clear_body_keep_sectpr(doc)
    if mode == "resume":
        _set_margins(doc, top=0.4, bottom=0.4, left=0.55, right=0.55)
    else:
        # Cover letter: roomier margins, since 1 page is the goal anyway.
        _set_margins(doc, top=0.7, bottom=0.7, left=0.8, right=0.8)

    if photo_path and photo_path.exists() and mode == "resume":
        _add_photo(doc, photo_path)

    md = mistune.create_markdown(renderer=None)
    tokens = md(markdown_text)

    seen_name = False
    contact_block_remaining = 0

    for tok in tokens:
        t = tok.get("type")

        if t == "heading":
            level = tok.get("attrs", {}).get("level", 2)
            text = _collect_text(tok.get("children", []))
            if level == 1 and not seen_name:
                _add_name(doc, text, mode=mode)
                seen_name = True
                contact_block_remaining = 4
            elif level <= 2:
                if mode == "cover":
                    # The only ## in a cover letter is the signature at the end —
                    # render as a left-aligned bold name, not a section heading.
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(4.0)
                    _add_run(p, text, bold=True, size_pt=11.0)
                else:
                    _add_section_heading(doc, text)
                contact_block_remaining = 0
            else:
                _add_subsection_heading(doc, text)
                contact_block_remaining = 0

        elif t == "paragraph":
            children = tok.get("children", [])
            if not _collect_text(children).strip():
                continue

            segments = _split_on_softbreaks(children)

            if contact_block_remaining > 0:
                for seg in segments:
                    if contact_block_remaining <= 0:
                        break
                    _add_subtitle(
                        doc,
                        seg,
                        is_first_under_name=(contact_block_remaining == 4),
                        mode=mode,
                    )
                    contact_block_remaining -= 1
            elif mode == "cover" and _collect_text(children).strip().upper().startswith("RE:"):
                # Subject line — left-aligned, bold (markdown already wraps it in **),
                # with extra space below to give visual "two enter" before the greeting.
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2.0)
                p.paragraph_format.space_after = Pt(18.0)
                _render_inline(p, children, base_size_pt=11.0)
            else:
                # For resumes, split lines that contain softbreaks (Technical Skills layout).
                # For cover letters, never split — softbreaks within a paragraph are rare
                # and should be treated as flowing text.
                if mode == "resume" and len(segments) > 1:
                    for seg in segments:
                        _add_body_paragraph(doc, seg, justify=False, mode=mode)
                else:
                    _add_body_paragraph(doc, children, mode=mode)

        elif t == "list":
            for item in tok.get("children", []):
                for child in item.get("children", []):
                    if child.get("type") in ("block_text", "paragraph"):
                        _add_bullet(doc, child.get("children", []), available_styles=available_styles)
                    elif child.get("type") == "list":
                        for sub_item in child.get("children", []):
                            for sub_child in sub_item.get("children", []):
                                if sub_child.get("type") in ("block_text", "paragraph"):
                                    _add_bullet(doc, sub_child.get("children", []), available_styles=available_styles)

        elif t == "thematic_break":
            doc.add_paragraph()

        elif t == "block_code":
            p = doc.add_paragraph()
            run = _add_run(p, tok.get("raw", ""))
            run.font.name = "Consolas"

        elif t == "blank_line":
            continue

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
